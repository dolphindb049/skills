import os
import signal
import queue
from docx import Document
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm

# ================= 配置区域 =================
API_KEY = "sk-MDQLUXvVuVESXdTSPfaH2PJE4uaXdC54YHpiXuKkTmwVxD8J" 
BASE_URL = "https://hk.n1n.ai/v1/"

INPUT_DIR = "docxfiles"       # 输入文件夹路径（如果是DOCX文档直接放这里，如果是PDF从上一步生成到这里）
OUTPUT_DIR = "finishfiles"     # 输出文件夹路径

MODEL = "gpt-5-nano" 
MAX_WORKERS = 4       # 并行处理文档数量

# ⭐ 目标语言设置： "en" 翻译为英文， "zh" 翻译为中文
TARGET_LANGUAGE = "zh" 

# ================= Global Interrupt Flag =================
stop_event = False

def signal_handler(sig, frame):
    global stop_event
    tqdm.write("\n🛑 检测到中断信号 (Ctrl+C). 正在等待任务安全停止...")
    stop_event = True

signal.signal(signal.SIGINT, signal_handler)

client = OpenAI(api_key=API_KEY, base_url=BASE_URL)

def has_chinese(text):
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            return True
    return False

def has_english_alpha(text):
    for char in text:
        if 'a' <= char.lower() <= 'z':
            return True
    return False

def needs_translation(text):
    """
    判断该段落是否需要翻译。
    翻译到英文时：看是否包含中文。
    翻译到中文时：看是否包含英文。
    """
    if TARGET_LANGUAGE == "en":
        return has_chinese(text)
    elif TARGET_LANGUAGE == "zh":
        return has_english_alpha(text)
    return True

def translate_text(text):
    """
    调用大模型进行纯翻译任务
    """
    if not text.strip() or text.isnumeric():
        return text

    sys_prompt = "请将此文本翻译为英文，直接输出翻译结果"
    if TARGET_LANGUAGE == "zh":
        sys_prompt = "请将此文本翻译为中文，直接输出翻译结果"

    try:
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": text}
            ],
            temperature=0.1,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return text

def process_single_docx(input_path, output_path, position_queue=None):
    """
    处理单个文档：逐段落/单元格提取并翻译，原地替换内容保存。
    """
    global stop_event
    if stop_event: return "🛑 已跳过"
    
    filename = os.path.basename(input_path)
    
    try:
        doc = Document(input_path)
    except Exception as e:
        return f"❌ {filename}: 读取文件失败 ({e})"

    tasks = []
    skipped_no_target = 0

    # --- 扫描正文 ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if needs_translation(text):
                tasks.append(para)
            else:
                skipped_no_target += 1

    # --- 扫描表格 ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        if needs_translation(text):
                            tasks.append(para)
                        else:
                            skipped_no_target += 1

    total = len(tasks)
    
    if total == 0:
        try:
            doc.save(output_path)
            return f"⚠️ {filename}: 无翻译内容，已直接复制保存"
        except Exception as e:
            return f"❌ {filename}: 保存失败 ({e})"

    pos = position_queue.get() if position_queue else 0
    try:
        desc_text = (filename[:15] + '..') if len(filename) > 15 else filename
        pbar = tqdm(total=total, desc=f"{desc_text:<17}", position=pos, leave=False, 
                    unit="段", bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}]")
        
        for para in tasks:
            if stop_event:
                pbar.close()
                return f"🛑 {filename}: 任务被中断"
            
            translated_content = translate_text(para.text.strip())
            # 直接替换原段落文本以保持样式
            para.text = translated_content
            
            pbar.update(1)
                
        pbar.close()

        if stop_event: return f"🛑 {filename}: 保存被取消"

        try:
            doc.save(output_path)
            return f"✅ {filename}: 完成"
        except PermissionError:
            return f"❌ {filename}: 保存失败！请先关闭已打开的 {output_path} 文件"
        except Exception as e:
            return f"❌ {filename}: 保存错误 ({e})"
    finally:
        if position_queue:
            position_queue.put(pos)

def process_directory(input_dir, output_dir):
    """
    批量翻译 DOCX 目录
    """
    global stop_event

    if not os.path.exists(input_dir):
        os.makedirs(input_dir)
        print(f"❌ 找不到输入文件夹 '{input_dir}'，已自动创建。请放入 DOCX 后重试。")
        return

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 查阅 docx 文件 (排除 ~ 开头的临时文件)
    docx_files = [f for f in os.listdir(input_dir) if f.endswith('.docx') and not f.startswith('~')]

    if not docx_files:
        print(f"⚠️ 文件夹 '{input_dir}' 中没有找到可翻译的 .docx 文件。")
        return

    print("=" * 45)
    print(f"📂 发现 {len(docx_files)} 个 DOCX 文档")
    print(f"🌐 目标语言: {'英文 (English)' if TARGET_LANGUAGE == 'en' else '中文 (Chinese)'}")
    print(f"⚡ 启用 {MAX_WORKERS} 个并发任务")
    print("👉 按 Ctrl+C 可以随时安全停止程序")
    print("=" * 45 + "\n")

    position_queue = queue.Queue()
    for i in range(MAX_WORKERS):
        position_queue.put(i)

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_to_file = {}
        for filename in docx_files:
            input_path = os.path.join(input_dir, filename)
            output_path = os.path.join(output_dir, filename)
            future = executor.submit(process_single_docx, input_path, output_path, position_queue)
            future_to_file[future] = filename
            
        try:
            for future in as_completed(future_to_file):
                if stop_event: break
                
                filename = future_to_file[future]
                try:
                    result_msg = future.result()
                    if result_msg:
                        tqdm.write(result_msg) 
                except Exception as e:
                    tqdm.write(f"❌ {filename} 异常: {e}")
                    
        except KeyboardInterrupt:
            stop_event = True
            print("\n🛑 主程序捕获中断...")

    if stop_event:
        print("\n🚫 翻译任务已停止。")
    else:
        print("\n🎉 所有翻译任务顺利结束！请前往 finishfiles 目录查看成果。")

if __name__ == "__main__":
    process_directory(INPUT_DIR, OUTPUT_DIR)
